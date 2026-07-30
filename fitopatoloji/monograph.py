# -*- coding: utf-8 -*-
"""Dijital Monografi Oluşturucu."""
from .common import *
from .rich_utils import to_reportlab, to_html
import webbrowser
try:
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


DEFAULT_FIELDS = [
    ("group_name", "Etmen grubu"), ("synonyms", "Sinonimler / eski adlar"),
    ("hosts", "Konukçular"), ("affected_organs", "Etkilenen organlar"),
    ("symptoms", "Belirtiler"), ("pathogen_features", "Etmenin özellikleri"),
    ("disease_cycle", "Hastalık döngüsü"), ("epidemiology", "Epidemiyoloji"),
    ("differential_diagnosis", "Ayırıcı teşhis"),
    ("distribution_turkey", "Türkiye dağılımı"), ("distribution_world", "Dünya dağılımı"),
    ("climate_notes", "İklim / çevre notları"),
    ("cultural_control", "Kültürel mücadele"), ("biological_control", "Biyolojik mücadele"),
    ("chemical_control", "Kimyasal mücadele / prensipler"),
    ("sources", "Kaynaklar"), ("notes", "Kişisel notlar"),
]


class MonographBuilder(tk.Toplevel):
    def __init__(self, master, database, paths, selected_ids=None):
        tk.Toplevel.__init__(self, master)
        self.db, self.paths = database, paths
        self.title("Dijital Monografi Oluşturucu")
        self.geometry("1120x720"); self.minsize(900, 620); self.transient(master)
        self.selected_ids = [int(x) for x in (selected_ids or [])]
        self.title_var = tk.StringVar(value="Fitopatoloji Monografisi")
        self.subtitle_var = tk.StringVar(value="Bitki hastalıkları bilimsel arşivi")
        self.author_var = tk.StringVar(value="")
        self.institution_var = tk.StringVar(value="")
        self.template_var = tk.StringVar(value="Bilimsel")
        self.photos_var = tk.BooleanVar(value=True)
        self.album_var = tk.BooleanVar(value=False)
        self.hide_empty_var = tk.BooleanVar(value=True)
        self.common_refs_var = tk.BooleanVar(value=True)
        self.status_var = tk.StringVar(value="Hastalıkları seçin ve sıralayın.")
        self.field_vars = {key: tk.BooleanVar(value=(key != "notes")) for key, _ in DEFAULT_FIELDS}
        self._build(); self._load_records(); self._load_projects()

    def _build(self):
        top = ttk.Frame(self, padding=12); top.pack(fill="x")
        ttk.Label(top, text="Dijital Monografi Oluşturucu", style="Title.TLabel").pack(side="left")
        ttk.Button(top, text="HTML önizleme", command=self.preview_html).pack(side="right", padx=3)
        ttk.Button(top, text="PDF oluştur", style="Primary.TButton", command=self.export_pdf).pack(side="right", padx=3)
        ttk.Button(top, text="DOCX oluştur", command=self.export_docx).pack(side="right", padx=3)
        ttk.Button(top, text="Taslağı kaydet", command=self.save_project).pack(side="right", padx=3)

        paned = ttk.Panedwindow(self, orient="horizontal"); paned.pack(fill="both", expand=True, padx=12)
        left = ttk.Frame(paned, padding=8); center = ttk.Frame(paned, padding=8); right = ttk.Frame(paned, padding=8)
        paned.add(left, weight=3); paned.add(center, weight=3); paned.add(right, weight=4)

        ttk.Label(left, text="1. Hastalıkları seç", style="Section.TLabel").pack(anchor="w")
        qrow=ttk.Frame(left); qrow.pack(fill="x", pady=6)
        self.query=tk.StringVar(); ttk.Entry(qrow,textvariable=self.query).pack(side="left",fill="x",expand=True)
        ttk.Button(qrow,text="Ara",command=self._load_records).pack(side="left",padx=(5,0))
        self.available=ttk.Treeview(left,columns=("name","group"),show="headings",selectmode="extended",height=16)
        self.available.heading("name",text="Hastalık / Etmen"); self.available.heading("group",text="Grup")
        self.available.column("name",width=260); self.available.column("group",width=110)
        self.available.pack(fill="both",expand=True)
        ttk.Button(left,text="Seçilenleri monografiye ekle →",command=self.add_selected).pack(fill="x",pady=6)
        ttk.Label(left,text="Kayıtlı taslaklar",style="Section.TLabel").pack(anchor="w",pady=(8,3))
        self.projects=ttk.Combobox(left,state="readonly"); self.projects.pack(fill="x")
        prow=ttk.Frame(left); prow.pack(fill="x",pady=4)
        ttk.Button(prow,text="Aç",command=self.load_project).pack(side="left",expand=True,fill="x")
        ttk.Button(prow,text="Sil",command=self.delete_project).pack(side="left",expand=True,fill="x",padx=(4,0))

        ttk.Label(center,text="2. Bölüm sırası",style="Section.TLabel").pack(anchor="w")
        self.chosen=ttk.Treeview(center,columns=("title",),show="headings",selectmode="extended")
        self.chosen.heading("title",text="Monografideki hastalıklar"); self.chosen.column("title",width=310)
        self.chosen.pack(fill="both",expand=True,pady=6)
        row=ttk.Frame(center); row.pack(fill="x")
        for text,cmd in [("↑ Yukarı",self.move_up),("↓ Aşağı",self.move_down),("Kaldır",self.remove_selected)]:
            ttk.Button(row,text=text,command=cmd).pack(side="left",expand=True,fill="x",padx=2)
        ttk.Label(center,text="Dahil edilecek bölümler",style="Section.TLabel").pack(anchor="w",pady=(12,4))
        canvas=tk.Canvas(center,height=185,highlightthickness=0); sb=ttk.Scrollbar(center,orient="vertical",command=canvas.yview)
        fields=ttk.Frame(canvas); fields.bind("<Configure>",lambda e:canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0,0),window=fields,anchor="nw"); canvas.configure(yscrollcommand=sb.set)
        canvas.pack(side="left",fill="both",expand=True); sb.pack(side="right",fill="y")
        for key,label in DEFAULT_FIELDS: ttk.Checkbutton(fields,text=label,variable=self.field_vars[key]).pack(anchor="w")

        ttk.Label(right,text="3. Kapak ve düzen",style="Section.TLabel").pack(anchor="w")
        form=ttk.Frame(right); form.pack(fill="x",pady=6)
        for label,var in [("Başlık",self.title_var),("Alt başlık",self.subtitle_var),("Yazar",self.author_var),("Kurum",self.institution_var)]:
            ttk.Label(form,text=label).pack(anchor="w",pady=(5,1)); ttk.Entry(form,textvariable=var).pack(fill="x")
        ttk.Label(form,text="Şablon").pack(anchor="w",pady=(8,1))
        ttk.Combobox(form,textvariable=self.template_var,values=("Bilimsel","Atlas","Ders Notu"),state="readonly").pack(fill="x")
        opts=ttk.LabelFrame(right,text="İçerik seçenekleri",padding=8); opts.pack(fill="x",pady=8)
        for text,var in [("Hastalık fotoğraflarını ekle",self.photos_var),("Sonda fotoğraf albümü oluştur",self.album_var),("Boş bölümleri gizle",self.hide_empty_var),("Ortak kaynakçayı tekilleştir",self.common_refs_var)]:
            ttk.Checkbutton(opts,text=text,variable=var).pack(anchor="w")
        ttk.Label(right,text="Monografi özeti",style="Section.TLabel").pack(anchor="w",pady=(8,3))
        self.summary=tk.Text(right,height=12,wrap="word",state="disabled"); self.summary.pack(fill="both",expand=True)
        ttk.Button(right,text="Özeti yenile",command=self.refresh_summary).pack(fill="x",pady=5)
        ttk.Label(self,textvariable=self.status_var,anchor="w",padding=(12,7)).pack(fill="x")

    def _load_records(self):
        if not hasattr(self,"available"): return
        self.available.delete(*self.available.get_children()); q=self.query.get().strip() if hasattr(self,"query") else ""
        for row in self.db.search(query=q):
            self.available.insert("","end",iid=str(row["id"]),values=("{} — {}".format(row["disease_name"],row["scientific_name"]),row["group_name"]))
        for did in self.selected_ids:
            if self.available.exists(str(did)): self.available.selection_add(str(did))
        if self.selected_ids: self.add_selected(); self.selected_ids=[]

    def add_selected(self):
        existing=set(self.chosen.get_children())
        for iid in self.available.selection():
            if iid not in existing:
                row=self.db.get(int(iid)); self.chosen.insert("","end",iid=iid,values=("{} — {}".format(row["disease_name"],row["scientific_name"]),))
        self.refresh_summary()
    def remove_selected(self):
        for iid in self.chosen.selection(): self.chosen.delete(iid)
        self.refresh_summary()
    def move_up(self):
        for iid in self.chosen.selection():
            pos=self.chosen.index(iid)
            if pos>0:self.chosen.move(iid,"",pos-1)
        self.refresh_summary()
    def move_down(self):
        for iid in reversed(self.chosen.selection()):
            pos=self.chosen.index(iid); total=len(self.chosen.get_children())
            if pos<total-1:self.chosen.move(iid,"",pos+1)
        self.refresh_summary()

    def _config(self):
        return {"title":self.title_var.get(),"subtitle":self.subtitle_var.get(),"author":self.author_var.get(),"institution":self.institution_var.get(),"template":self.template_var.get(),"photos":self.photos_var.get(),"album":self.album_var.get(),"hide_empty":self.hide_empty_var.get(),"common_refs":self.common_refs_var.get(),"fields":[k for k,_ in DEFAULT_FIELDS if self.field_vars[k].get()],"disease_ids":[int(x) for x in self.chosen.get_children()]}
    def refresh_summary(self):
        cfg=self._config(); photos=sum(len(self.db.image_attachments(i)) for i in cfg["disease_ids"])
        refs=sum(len(self.db.references(i)) for i in cfg["disease_ids"])
        text="{}\n\n{} hastalık • {} fotoğraf • {} yapılandırılmış kaynak\nŞablon: {}\nBölüm sayısı: {}\n\nPDF ve HTML çıktısı kapak, içindekiler, hastalık bölümleri ve indeks içerir.".format(cfg["title"],len(cfg["disease_ids"]),photos,refs,cfg["template"],len(cfg["fields"]))
        self.summary.configure(state="normal"); self.summary.delete("1.0","end"); self.summary.insert("1.0",text); self.summary.configure(state="disabled")

    def save_project(self):
        cfg=self._config()
        if not cfg["disease_ids"]: messagebox.showinfo(APP_NAME,"Önce en az bir hastalık ekleyin.",parent=self); return
        name=simpledialog.askstring("Monografi taslağı","Taslak adı:",initialvalue=cfg["title"],parent=self)
        if name: self.db.save_monograph_project(name,cfg); self._load_projects(); self.status_var.set("Taslak kaydedildi: "+name)
    def _load_projects(self):
        self._project_rows=self.db.monograph_projects(); self.projects["values"]=[r["name"] for r in self._project_rows]
        if self._project_rows:self.projects.current(0)
    def load_project(self):
        idx=self.projects.current()
        if idx<0:return
        cfg=json.loads(self._project_rows[idx]["config_json"])
        for attr,key in [(self.title_var,"title"),(self.subtitle_var,"subtitle"),(self.author_var,"author"),(self.institution_var,"institution"),(self.template_var,"template")]: attr.set(cfg.get(key,""))
        for var,key in [(self.photos_var,"photos"),(self.album_var,"album"),(self.hide_empty_var,"hide_empty"),(self.common_refs_var,"common_refs")]:var.set(bool(cfg.get(key)))
        selected=set(cfg.get("fields",[]))
        for k in self.field_vars:self.field_vars[k].set(k in selected)
        self.chosen.delete(*self.chosen.get_children())
        for did in cfg.get("disease_ids",[]):
            row=self.db.get(did)
            if row:self.chosen.insert("","end",iid=str(did),values=("{} — {}".format(row["disease_name"],row["scientific_name"]),))
        self.refresh_summary(); self.status_var.set("Taslak açıldı.")
    def delete_project(self):
        idx=self.projects.current()
        if idx>=0 and messagebox.askyesno(APP_NAME,"Seçili taslak silinsin mi?",parent=self): self.db.delete_monograph_project(self._project_rows[idx]["id"]); self._load_projects()

    def _html(self,cfg):
        toc=[]; chapters=[]; all_refs=[]; hosts=set(); agents=set()
        for no,did in enumerate(cfg["disease_ids"],1):
            r=self.db.get(did); rich=self.db.rich_text(did); toc.append('<li><a href="#d{}">{}. {}</a></li>'.format(did,no,xml_escape(r["disease_name"])))
            agents.add((r["scientific_name"] or "").strip());
            for h in re.split(r"[,;\n]",r["hosts"] or ""):
                if h.strip():hosts.add(h.strip())
            body=['<section class="chapter" id="d{}"><h1>{}. {}</h1><h2><em>{}</em></h2>'.format(did,no,xml_escape(r["disease_name"]),xml_escape(r["scientific_name"]))]
            for key,label in DEFAULT_FIELDS:
                if key not in cfg["fields"]:continue
                value=(r[key] or "").strip()
                if not value and cfg["hide_empty"]:continue
                body.append('<h3>{}</h3><div>{}</div>'.format(label,to_html(value,rich.get(key,{}))))
            refs=self.db.references(did); all_refs.extend([x["citation"] for x in refs if x["citation"]])
            linked=self.db.disease_literature(did)
            for x in linked:
                citation="{} ({}) {}. {}{}".format(x["authors"],x["year_text"],x["title"],x["journal"],(" DOI: "+x["doi"]) if x["doi"] else "").strip()
                if citation: all_refs.append(citation)
            if linked:
                body.append('<h3>Yapılandırılmış literatür</h3><ol>'+''.join('<li>{}</li>'.format(xml_escape("{} ({}) {}".format(x["authors"],x["year_text"],x["title"]))) for x in linked)+'</ol>')
            if cfg["photos"]:
                imgs=[]
                for p in self.db.image_attachments(did):
                    path=os.path.abspath(os.path.join(self.paths.base,p["relative_path"])); cap=((p["image_category"] or "Genel")+" — "+(p["title"] or p["description"] or os.path.basename(path)))
                    if os.path.isfile(path):imgs.append('<figure><img src="file:///{}"><figcaption>{}</figcaption></figure>'.format(path.replace('\\','/'),xml_escape(cap)))
                if imgs:body.append('<h3>Fotoğraflar</h3><div class="gallery">'+''.join(imgs)+'</div>')
            body.append('</section>'); chapters.append(''.join(body))
        refs=sorted(set(x.strip() for x in all_refs if x.strip())) if cfg["common_refs"] else all_refs
        refhtml=''.join('<li>{}</li>'.format(xml_escape(x)) for x in refs)
        index='<h1>Bilimsel İndeks</h1><h2>Etmenler</h2><p>{}</p><h2>Konukçular</h2><p>{}</p>'.format('<br>'.join(sorted(agents,key=str.lower)),'<br>'.join(sorted(hosts,key=str.lower)))
        css='body{font-family:Segoe UI,Arial,sans-serif;color:#24332c;margin:0;background:#eef3f0}.page{max-width:900px;margin:24px auto;background:white;padding:55px;box-shadow:0 3px 16px #999}.cover{text-align:center;min-height:700px;display:flex;flex-direction:column;justify-content:center}.chapter{page-break-before:always}.chapter h1{border-bottom:3px solid #517b63;padding-bottom:8px}.gallery{display:grid;grid-template-columns:repeat(2,1fr);gap:16px}.gallery img{width:100%;height:240px;object-fit:contain;background:#f2f2f2}figure{margin:0;border:1px solid #ddd;padding:8px}figcaption{text-align:center;font-size:12px}.toc a{color:#365d46;text-decoration:none}h1,h2,h3{color:#365d46}@media print{body{background:white}.page{box-shadow:none;margin:0;max-width:none}}'
        return '<!doctype html><html><head><meta charset="utf-8"><title>{}</title><style>{}</style></head><body><main class="page"><section class="cover"><h1>{}</h1><h2>{}</h2><p>{}</p><p>{}</p><p>{}</p></section><section class="chapter toc"><h1>İçindekiler</h1><ol>{}</ol></section>{}<section class="chapter"><h1>Ortak Kaynakça</h1><ol>{}</ol></section><section class="chapter">{}</section></main></body></html>'.format(xml_escape(cfg["title"]),css,xml_escape(cfg["title"]),xml_escape(cfg["subtitle"]),xml_escape(cfg["author"]),xml_escape(cfg["institution"]),dt.datetime.now().strftime("%d.%m.%Y"),''.join(toc),''.join(chapters),refhtml,index)

    def preview_html(self):
        cfg=self._config()
        if not cfg["disease_ids"]:messagebox.showinfo(APP_NAME,"Önce hastalık seçin.",parent=self);return
        path=os.path.join(self.paths.exports,"Monografi_Onizleme.html")
        with open(path,"w",encoding="utf-8") as f:f.write(self._html(cfg))
        webbrowser.open("file:///"+path.replace('\\','/')); self.status_var.set("HTML önizleme açıldı.")

    def export_docx(self):
        cfg=self._config()
        if not cfg["disease_ids"]:
            messagebox.showinfo(APP_NAME,"Önce hastalık seçin.",parent=self); return
        if not DOCX_AVAILABLE:
            messagebox.showerror(APP_NAME,"DOCX için python-docx bulunamadı.",parent=self); return
        out=filedialog.asksaveasfilename(parent=self,title="Monografi DOCX",initialdir=self.paths.exports,initialfile=re.sub(r'[^A-Za-z0-9_-]+','_',cfg['title']).strip('_')+'.docx',defaultextension='.docx',filetypes=[('Word belgesi','*.docx')])
        if not out:return
        try:
            doc=Document(); sec=doc.sections[0]; sec.top_margin=Cm(1.8); sec.bottom_margin=Cm(1.8); sec.left_margin=Cm(2); sec.right_margin=Cm(2)
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(cfg['title']); r.bold=True; r.font.size=Pt(24)
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(cfg['subtitle']).italic=True
            for value in (cfg['author'],cfg['institution'],dt.datetime.now().strftime('%d.%m.%Y')):
                p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(value)
            doc.add_page_break(); doc.add_heading('İçindekiler',level=1)
            for no,did in enumerate(cfg['disease_ids'],1):
                row=self.db.get(did); doc.add_paragraph('{}. {}'.format(no,row['disease_name']),style='List Number')
            all_refs=[]
            for no,did in enumerate(cfg['disease_ids'],1):
                row=self.db.get(did); doc.add_page_break(); doc.add_heading('{}. {}'.format(no,row['disease_name']),level=1)
                p=doc.add_paragraph(); rr=p.add_run(row['scientific_name'] or ''); rr.italic=True
                for key,label in DEFAULT_FIELDS:
                    if key not in cfg['fields']:continue
                    value=(row[key] or '').strip()
                    if not value and cfg['hide_empty']:continue
                    doc.add_heading(label,level=2); doc.add_paragraph(value)
                linked=self.db.disease_literature(did)
                if linked:
                    doc.add_heading('Yapılandırılmış literatür',level=2)
                    for x in linked:
                        citation="{} ({}) {}. {}{}".format(x['authors'],x['year_text'],x['title'],x['journal'],(' DOI: '+x['doi']) if x['doi'] else '').strip(); all_refs.append(citation); doc.add_paragraph(citation,style='List Bullet')
                if cfg['photos']:
                    photos=self.db.image_attachments(did)
                    if photos: doc.add_heading('Fotoğraflar',level=2)
                    for ph in photos:
                        path=os.path.join(self.paths.base,ph['relative_path'])
                        if os.path.isfile(path):
                            try: doc.add_picture(path,width=Cm(13.5))
                            except Exception: pass
                            cap='{} — {}'.format(ph['image_category'] or 'Genel',ph['title'] or ph['description'] or os.path.basename(path)); cp=doc.add_paragraph(cap); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                            rights=' • '.join(x for x in (ph['photographer'],ph['copyright_owner'],ph['license_text'],ph['scale_info'],ph['location_text']) if x)
                            if rights:
                                rp=doc.add_paragraph(rights); rp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            if all_refs:
                doc.add_page_break(); doc.add_heading('Ortak Kaynakça',level=1)
                refs=sorted(set(all_refs),key=str.lower) if cfg['common_refs'] else all_refs
                for ref in refs: doc.add_paragraph(ref,style='List Number')
            doc.save(out); self.db.save_monograph_export(cfg['title'],out,len(cfg['disease_ids'])); messagebox.showinfo(APP_NAME,'DOCX oluşturuldu:\n'+out,parent=self)
        except Exception as exc:
            messagebox.showerror(APP_NAME,'DOCX oluşturulamadı:\n{}'.format(exc),parent=self)

    def export_pdf(self):
        cfg=self._config()
        if not cfg["disease_ids"]:messagebox.showinfo(APP_NAME,"Önce hastalık seçin.",parent=self);return
        if not REPORTLAB_AVAILABLE:messagebox.showerror(APP_NAME,"PDF için ReportLab bulunamadı.",parent=self);return
        out=filedialog.asksaveasfilename(parent=self,title="Monografi PDF",initialdir=self.paths.exports,initialfile=re.sub(r'[^A-Za-z0-9_-]+','_',cfg['title']).strip('_')+'.pdf',defaultextension='.pdf',filetypes=[('PDF','*.pdf')])
        if not out:return
        try:
            styles=getSampleStyleSheet(); normal=styles['BodyText']; normal.fontName='Helvetica'; normal.fontSize=9; normal.leading=13
            title=ParagraphStyle('MonoTitle',parent=styles['Title'],fontSize=24,leading=30,alignment=TA_CENTER,spaceAfter=16)
            h1=ParagraphStyle('MonoH1',parent=styles['Heading1'],fontSize=17,leading=21,spaceAfter=10)
            h2=ParagraphStyle('MonoH2',parent=styles['Heading2'],fontSize=11,leading=14,spaceBefore=8,spaceAfter=4)
            story=[Spacer(1,5*cm),Paragraph(xml_escape(cfg['title']),title),Paragraph(xml_escape(cfg['subtitle']),ParagraphStyle('sub',parent=normal,fontSize=13,alignment=TA_CENTER)),Spacer(1,1*cm),Paragraph(xml_escape(cfg['author']),ParagraphStyle('a',parent=normal,alignment=TA_CENTER)),Paragraph(xml_escape(cfg['institution']),ParagraphStyle('i',parent=normal,alignment=TA_CENTER)),PageBreak(),Paragraph('İçindekiler',h1)]
            rows=[]
            for no,did in enumerate(cfg['disease_ids'],1):
                r=self.db.get(did); rows.append([str(no),Paragraph(xml_escape(r['disease_name']),normal)])
            story.append(Table(rows,colWidths=[1*cm,15*cm])); all_refs=[]; agents=set(); hosts=set()
            for no,did in enumerate(cfg['disease_ids'],1):
                r=self.db.get(did); rich=self.db.rich_text(did); story.extend([PageBreak(),Paragraph('{}. {}'.format(no,xml_escape(r['disease_name'])),h1),Paragraph('<i>{}</i>'.format(xml_escape(r['scientific_name'])),h2)])
                agents.add((r['scientific_name'] or '').strip())
                for host in re.split(r'[,;\n]',r['hosts'] or ''):
                    if host.strip():hosts.add(host.strip())
                for key,label in DEFAULT_FIELDS:
                    if key not in cfg['fields']:continue
                    value=(r[key] or '').strip()
                    if not value and cfg['hide_empty']:continue
                    story.extend([Paragraph(label,h2),Paragraph(to_reportlab(value,rich.get(key,{})),normal)])
                refs=self.db.references(did); all_refs.extend([x['citation'] for x in refs if x['citation']])
                linked=self.db.disease_literature(did)
                if linked:
                    story.append(Paragraph('Yapılandırılmış literatür',h2))
                    for x in linked:
                        citation="{} ({}) {}. {}{}".format(x['authors'],x['year_text'],x['title'],x['journal'],(' DOI: '+x['doi']) if x['doi'] else '').strip()
                        all_refs.append(citation); story.append(Paragraph('• '+xml_escape(citation),normal))
                if cfg['photos']:
                    cells=[]
                    for p in self.db.image_attachments(did):
                        path=os.path.join(self.paths.base,p['relative_path'])
                        if os.path.isfile(path):
                            try:
                                im=RLImage(path); im._restrictSize(7.2*cm,5*cm); cap=Paragraph(xml_escape((p['image_category'] or 'Genel')+' — '+(p['title'] or p['description'] or os.path.basename(path))),ParagraphStyle('cap',parent=normal,fontSize=7,alignment=TA_CENTER)); cells.append([im,cap])
                            except Exception:pass
                    if cells:
                        grid=[]
                        for i in range(0,len(cells),2):grid.append(cells[i:i+2]+([''] if len(cells[i:i+2])==1 else []))
                        story.extend([Paragraph('Fotoğraflar',h2),Table(grid,colWidths=[8*cm,8*cm],style=[('VALIGN',(0,0),(-1,-1),'TOP'),('BOX',(0,0),(-1,-1),.25,colors.lightgrey),('INNERGRID',(0,0),(-1,-1),.25,colors.lightgrey)])])
            refs=sorted(set(x.strip() for x in all_refs if x.strip())) if cfg['common_refs'] else all_refs
            story.extend([PageBreak(),Paragraph('Ortak Kaynakça',h1)])
            for i,x in enumerate(refs,1):story.append(Paragraph('{}. {}'.format(i,xml_escape(x)),normal))
            story.extend([PageBreak(),Paragraph('Bilimsel İndeks',h1),Paragraph('Etmenler',h2),Paragraph('<br/>'.join(xml_escape(x) for x in sorted(agents,key=str.lower)),normal),Paragraph('Konukçular',h2),Paragraph('<br/>'.join(xml_escape(x) for x in sorted(hosts,key=str.lower)),normal)])
            def footer(canvas,doc):
                canvas.saveState(); canvas.setFont('Helvetica',8); canvas.drawCentredString(A4[0]/2,0.7*cm,str(doc.page)); canvas.restoreState()
            SimpleDocTemplate(out,pagesize=A4,rightMargin=1.5*cm,leftMargin=1.5*cm,topMargin=1.5*cm,bottomMargin=1.3*cm,title=cfg['title'],author=cfg['author'] or APP_NAME).build(story,onFirstPage=footer,onLaterPages=footer)
            self.db.save_monograph_export(cfg['title'],out,len(cfg['disease_ids'])); messagebox.showinfo(APP_NAME,"Monografi oluşturuldu:\n"+out,parent=self)
        except Exception as exc:messagebox.showerror(APP_NAME,"Monografi oluşturulamadı:\n{}".format(exc),parent=self)
